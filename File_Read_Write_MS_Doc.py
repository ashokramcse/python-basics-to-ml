# -*- coding: utf-8 -*-
"""
Created on Wed Jan 17 03:09:05 2018

"""

from docx import Document
from docx.shared import Inches

document = Document('C:\\Users\\RPS\\Desktop\\DocumentOne.docx')
for para in document.paragraphs:
    print(para.text)

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture('butterfly.jpg', width=Inches(1.25))

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
recordset = [ [1, "101", "Spam"], [2, "42", "Eggs"], [3, "631", "Spam, spam, eggs, and spam"]]
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item[0])
    row_cells[1].text = str(item[1])
    row_cells[2].text = item[2]

document.add_page_break()

document.save('demo.docx')